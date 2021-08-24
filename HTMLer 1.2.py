import docx
import pyperclip

doc = docx.Document('Ready.docx')
start_div = "<div style=\"width: 300px; display: inline-block; vertical-align: top;; padding: 5px; margin: 5px;\">"
start_p_img = "<img width=\"300\" height=\"400\" alt=\"\" src=\"https://mkala.rpa-mu.ru/Media/mkala/dostijeniya_studentov/2020-2021/zamena\" />"
start_p_style = "<p style=\"text-align: center;\">"
start_p = "<p>"
end_p = "</p>"
end_div = "</div>"
start_strong = "<strong>"
end_strong = "</strong>"

start_stop = []
final_text = ""


for i in range(len(doc.paragraphs)):
    if doc.paragraphs[i].text == "":
        start_stop.append(i)

start = 0

for st_st in start_stop:
    stop = st_st

    if "." in doc.paragraphs[start].text[:len(doc.paragraphs[start].text) - 5]:
        tochka = doc.paragraphs[start].text.find(".")
        doc.paragraphs[start].text = doc.paragraphs[start].text[tochka + 1:]
        if doc.paragraphs[start].text[0] == " ":
            doc.paragraphs[start].text = doc.paragraphs[start].text[1:]
            doc.paragraphs[start].text = doc.paragraphs[start].text.title()
        if doc.paragraphs[start].text != "":
            if doc.paragraphs[start].text[len(doc.paragraphs[start].text) - 1] == " ":
                doc.paragraphs[start].text = doc.paragraphs[start].text[:len(doc.paragraphs[start].text)-1]
        
    photo_studenta = start_p_img.replace("zamena", doc.paragraphs[start].text.replace(" ", "%20") + ".jpg")
    
    div_block = ""
    div_block = start_div + div_block + "\n"
    div_block = div_block + start_p_style + photo_studenta + end_p + "\n"
    div_block = div_block + start_p_style + start_strong + doc.paragraphs[start].text + end_strong + end_p + "\n"

    if doc.paragraphs[start+1].text == "(Юридический колледж)":
        div_block = div_block + start_p_style + start_strong + doc.paragraphs[start+1].text + end_strong + end_p + "\n"
        div_block = div_block + start_p_style + start_strong + doc.paragraphs[start+2].text + end_strong + end_p + "\n"

        for texting in range(start + 3, stop):
            div_block = div_block + start_p + doc.paragraphs[texting].text + end_p + "\n"

    else:
        div_block = div_block + start_p_style + start_strong + doc.paragraphs[start+1].text + end_strong + end_p + "\n"

        for texting in range(start + 2, stop):
            div_block = div_block + start_p + doc.paragraphs[texting].text + end_p + "\n"
        
    div_block = div_block + end_div + "\n" + "\n"
    
    final_text = final_text + div_block
    
    start = st_st + 1

final_text = "<div style=\"width: 1000px;\">" + final_text + end_div
pyperclip.copy(final_text)
print("READY!")
