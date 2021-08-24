import os
import random
import docx
from docx import Document

directory = tuple()
directory = (os.listdir(str(input())))
length = len(directory)
for i in range (1, len(directory)):
    file_name = directory[i]
    file_name = file_name.replace('.docx', '')
    doc = Document('Shablon.docx')
    kolvo_tablic = len(doc.tables)
    kolvo_abzacev = len(doc.paragraphs)
    for m in range (0, kolvo_tablic):
        doc.tables[m].rows[1].cells[2].text = str(random.randint(53, 81))
        doc.tables[m].rows[1].cells[2].paragraphs[0].style = 'Yacheika' 
    ii = 17
    while ii < kolvo_abzacev - 17:
        doc.paragraphs[ii].runs[len(doc.paragraphs[ii].runs)-1].add_break(docx.enum.text.WD_BREAK.PAGE) 
        ii = ii + 18
    ii = 2
    while ii < kolvo_abzacev:
        doc.paragraphs[ii].text = doc.paragraphs[ii].text.replace('Иванов Иван Иванович', file_name)
        doc.paragraphs[ii].style = 'FIO'
        ii = ii + 18
    doc.paragraphs[kolvo_abzacev - 1].text = ''
    konechnoe_imya_faila = file_name+".docx"
    doc.save(konechnoe_imya_faila)
